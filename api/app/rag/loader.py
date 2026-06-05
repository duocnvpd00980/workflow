"""DocumentLoaderService v2.4 — Sửa lỗi luồng Metadata Tagging."""

from __future__ import annotations

import asyncio
import logging
import re
from dataclasses import dataclass, field
from pathlib import Path
from typing import Dict, Iterator, List, Optional

logger = logging.getLogger(__name__)


@dataclass
class LoadedDocument:
    text: str
    metadata: Dict = field(default_factory=dict)


class DocumentLoader:
    """
    Hỗ trợ: web (crawl4ai), PDF, DOCX, XLSX, TXT, CSV, Markdown.
    """

    SUPPORTED_EXTENSIONS: List[str] = [".pdf", ".docx", ".txt", ".csv", ".xlsx", ".md"]
    MAX_FILE_SIZE_MB: int = 50

    _SKIP_PATTERNS = ("/.git/", "/.venv/", "/venv/", "/node_modules/", "/__pycache__/", "/dist/", "/build/")
    _SKIP_PREFIXES = (".", "~")
    _EXCESS_NEWLINE_RE: re.Pattern = re.compile(r"\n{3,}")

    def __init__(self, max_retries: int = 3) -> None:
        self._max_retries = max_retries

    # ── Text Cleaning ─────────────────────────────────────────────────────────

    def _clean_text(self, raw: str) -> str:
        if not raw:
            return ""
        lines = [line.strip() for line in raw.split("\n")]
        return self._EXCESS_NEWLINE_RE.sub("\n\n", "\n".join(lines)).strip()

    # ── Web Loading (crawl4ai) ────────────────────────────────────────────────

    def load_web(self, url: str, document_type: str = "web_page") -> LoadedDocument:
        """
        Crawl URL bằng crawl4ai — hỗ trợ JS-rendered page, SPA, lazy-load.
        """
        try:
            loop = asyncio.get_running_loop()
            import concurrent.futures
            with concurrent.futures.ThreadPoolExecutor(max_workers=1) as pool:
                # FIX 1: Truyền chuẩn document_type vào coroutine
                future = pool.submit(asyncio.run, self._crawl4ai(url, document_type))
                return future.result()
        except RuntimeError:
            return asyncio.run(self._crawl4ai(url, document_type))

    async def _crawl4ai(self, url: str, document_type: str) -> LoadedDocument:
        """FIX 1: Nhận thêm tham số document_type từ hàm gọi."""
        try:
            from crawl4ai import AsyncWebCrawler, BrowserConfig, CrawlerRunConfig
            from crawl4ai.content_filter_strategy import PruningContentFilter
            from crawl4ai.markdown_generation_strategy import DefaultMarkdownGenerator
        except ImportError:
            raise ImportError("Cài đặt: pip install crawl4ai && crawl4ai-setup")

        browser_cfg = BrowserConfig(headless=True, verbose=False)
        md_generator = DefaultMarkdownGenerator(
            content_filter=PruningContentFilter(threshold=0.48, threshold_type="fixed"),
        )
        run_cfg = CrawlerRunConfig(
            markdown_generator=md_generator,
            wait_until="networkidle",
            page_timeout=30_000,
            cache_mode="bypass",
        )

        last_exc: Exception | None = None
        for attempt in range(self._max_retries):
            try:
                async with AsyncWebCrawler(config=browser_cfg) as crawler:
                    result = await crawler.arun(url=url, config=run_cfg)

                if not result.success:
                    raise ValueError(f"crawl4ai thất bại: {result.error_message}")

                text = (
                    result.markdown.fit_markdown
                    or result.markdown.raw_markdown
                    or ""
                ).strip()

                if not text:
                    raise ValueError("Không trích xuất được nội dung từ URL.")

                return LoadedDocument(
                    text=self._clean_text(text),
                    metadata={
                        "source_id":      url,
                        "source_url":     url,
                        "document_type":  document_type, # FIX 2: Không gán cứng "web_page" nữa
                        "crawl_engine":   "crawl4ai",
                    },
                )

            except Exception as exc:
                last_exc = exc
                logger.warning("[Loader] crawl4ai lần %d/%d thất bại: %s — %s",
                               attempt + 1, self._max_retries, url, exc)
                if attempt < self._max_retries - 1:
                    await asyncio.sleep(2 ** attempt)

        raise ValueError(
            f"Không thể crawl URL sau {self._max_retries} lần thử: {url}"
        ) from last_exc

    # ── File Reading ──────────────────────────────────────────────────────────

    def _check_file_size(self, path: Path) -> None:
        size_mb = path.stat().st_size / (1024 * 1024)
        if size_mb > self.MAX_FILE_SIZE_MB:
            raise ValueError(f"File {path.name} ({size_mb:.1f} MB) vượt giới hạn {self.MAX_FILE_SIZE_MB} MB")

    def _read_pdf(self, path: Path) -> str:
        try:
            import fitz
        except ImportError:
            raise ImportError("uv add pymupdf")

        doc = fitz.open(str(path))
        pages = []
        for page in doc:
            text = page.get_text("text")
            if text.strip():
                pages.append(text)
                continue
            blocks = page.get_text("blocks")
            block_text = "\n".join(b[4] for b in blocks if b[4].strip())
            if block_text.strip():
                pages.append(block_text)
                continue
            try:
                import io
                import pytesseract
                from PIL import Image
                pix = page.get_pixmap(dpi=200)
                img = Image.open(io.BytesIO(pix.tobytes("png")))
                ocr_text = pytesseract.image_to_string(img, lang="vie+eng")
                if ocr_text.strip():
                    pages.append(ocr_text)
            except Exception as e:
                logger.warning("[Loader] OCR page %d thất bại: %s", page.number + 1, e)

        if not pages:
            raise ValueError(f"Không đọc được nội dung (kể cả OCR): {path.name}")
        return "\n\n".join(pages)

    def _read_docx(self, path: Path) -> str:
        try:
            from docx import Document
        except ImportError:
            raise ImportError("pip install python-docx")
        doc = Document(str(path))
        return "\n".join(p.text for p in doc.paragraphs if p.text.strip())

    def _read_xlsx(self, path: Path) -> str:
        try:
            from openpyxl import load_workbook
        except ImportError:
            raise ImportError("pip install openpyxl")
        wb = load_workbook(str(path), data_only=True)
        rows: List[str] = []
        for sheet in wb.worksheets:
            for row in sheet.iter_rows(values_only=True):
                row_text = " | ".join(str(c) for c in row if c is not None)
                if row_text.strip():
                    rows.append(row_text)
        return "\n".join(rows)

    def _read_txt(self, path: Path) -> str:
        for enc in ("utf-8", "utf-16", "cp1252", "iso-8859-1"):
            try:
                return path.read_text(encoding=enc)
            except UnicodeDecodeError:
                continue
        raise ValueError(f"Không thể decode file: {path}")

    _READERS = {
        ".pdf":  _read_pdf,
        ".docx": _read_docx,
        ".xlsx": _read_xlsx,
        ".txt":  _read_txt,
        ".csv":  _read_txt,
        ".md":   _read_txt,
    }

    def _read_file(self, path: Path) -> str:
        reader = self._READERS.get(path.suffix.lower())
        if not reader:
            raise ValueError(f"Không có reader cho định dạng: {path.suffix}")
        return reader(self, path)

    # ── Single File ───────────────────────────────────────────────────────────

    def load_file(self, path: str, document_type: str = "product_knowledge") -> List[LoadedDocument]:
        file_path = Path(path)
        if not file_path.exists():
            raise FileNotFoundError(f"Không tìm thấy file: {path}")
        self._check_file_size(file_path)
        cleaned = self._clean_text(self._read_file(file_path))
        if not cleaned:
            raise ValueError(f"Không trích xuất được nội dung từ: {file_path.name}")
        
        # FIX 3: Ghi đè chính xác tag được chỉ định qua tham số đầu vào
        metadata = self._file_metadata(file_path, document_type=document_type)
        return [LoadedDocument(text=cleaned, metadata=metadata)]

    def _file_metadata(self, path: Path, document_type: str = "product_knowledge") -> Dict:
        """Cập nhật để hàm tạo metadata nhận diện được type động."""
        return {
            "source_id":        path.name,
            "file_name":        path.name,
            "file_extension":   path.suffix.lower().lstrip("."),
            "file_size_bytes":  path.stat().st_size,
            "absolute_path":    str(path.resolve()),
            "document_type":    document_type, # Nhận động từ ngoài
        }

    # ── Directory ─────────────────────────────────────────────────────────────

    def _should_skip(self, path: Path) -> bool:
        posix = path.as_posix()
        if any(pat in posix for pat in self._SKIP_PATTERNS):
            return True
        if any(part.startswith(self._SKIP_PREFIXES) for part in path.parts):
            return True
        return False

    def _iter_files(self, dir_path: Path) -> Iterator[Path]:
        for file_path in dir_path.rglob("*"):
            if (file_path.is_file()
                    and not self._should_skip(file_path)
                    and file_path.suffix.lower() in self.SUPPORTED_EXTENSIONS):
                yield file_path

    def load_directory(self, path: str, document_type: str = "product_knowledge") -> List[LoadedDocument]:
        """FIX 4: Thêm tham số document_type khi đọc hàng loạt thư mục."""
        dir_path = Path(path)
        if not dir_path.is_dir():
            raise NotADirectoryError(f"Không phải thư mục: {path}")
        docs: List[LoadedDocument] = []
        for file_path in self._iter_files(dir_path):
            try:
                self._check_file_size(file_path)
                cleaned = self._clean_text(self._read_file(file_path))
                if cleaned:
                    # FIX 5: Ép tag động cho từng file trong directory
                    meta = self._file_metadata(file_path, document_type=document_type)
                    docs.append(LoadedDocument(text=cleaned, metadata=meta))
            except Exception as exc:
                logger.warning("[Loader] Bỏ qua %s: %s", file_path, exc)
        logger.info("[Loader] Nạp %d tài liệu từ %s với tag: %s", len(docs), path, document_type)
        return docs
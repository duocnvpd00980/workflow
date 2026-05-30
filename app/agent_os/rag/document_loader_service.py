"""DocumentLoaderService v2.1 — Tải và làm sạch tài liệu cho RAG ZERO v2.1.
Không phụ thuộc llama-index. Trả về plain dict để pipeline RAG xử lý.
"""

from __future__ import annotations
import re
import logging
from pathlib import Path
from dataclasses import dataclass, field
from typing import List, Dict, Iterator
from time import sleep

import trafilatura

logger = logging.getLogger(__name__)


@dataclass
class LoadedDocument:
    """Thay thế llama_index.Document — chỉ cần text + metadata."""
    text: str
    metadata: Dict = field(default_factory=dict)

    
class DocumentLoader:
    """
    Dịch vụ tải tài liệu chuẩn hóa cho RAG ZERO v2.1.
    Không phụ thuộc llama-index. Xử lý tiền kỳ dữ liệu thô.
    """

    SUPPORTED_EXTENSIONS: List[str] = [".pdf", ".docx", ".txt", ".csv", ".xlsx", ".md"]
    MAX_FILE_SIZE_MB: int = 50

    _EXCESS_NEWLINE_PATTERN: re.Pattern = re.compile(r"\n{3,}")

    def __init__(self, request_timeout_seconds: int = 15, max_retries: int = 3) -> None:
        self._timeout = request_timeout_seconds
        self._max_retries = max_retries

    def _clean_text(self, raw_text: str) -> str:
        """Làm sạch văn bản thô, giữ ranh giới đoạn văn."""
        if not raw_text:
            return ""
        lines = [line.strip() for line in raw_text.split("\n")]
        joined = "\n".join(lines)
        cleaned = self._EXCESS_NEWLINE_PATTERN.sub("\n\n", joined)
        return cleaned.strip()

    # ── Web Loading ───────────────────────────────────────────────────────────

    def load_web(self, url: str) -> LoadedDocument:
        """Cào web với retry + backoff."""
        last_exc = None
        for attempt in range(self._max_retries):
            try:
                downloaded = trafilatura.fetch_url(url, timeout=self._timeout)
                if downloaded:
                    break
            except Exception as exc:
                last_exc = exc
                if attempt < self._max_retries - 1:
                    sleep(2 ** attempt)  # exponential backoff: 1s, 2s, 4s
                continue
        else:
            raise ValueError(f"Không thể tải URL sau {self._max_retries} lần thử: {url}") from last_exc

        extracted = trafilatura.extract(
            downloaded,
            include_comments=False,
            include_tables=True,
            no_fallback=False,
            target_language="vi",  # ← tối ưu: chỉ extract tiếng Việt
        )

        if not extracted:
            raise ValueError(f"Không trích xuất được nội dung từ URL: {url}")

        return LoadedDocument(
            text=self._clean_text(extracted),
            metadata={
                "source_url": url,
                "document_type": "web_page",
                "storage_layer": "external_crawl",
            },
        )

    # ── File Loading ──────────────────────────────────────────────────────────

    def _check_file_size(self, file_path: Path) -> None:
        """Giới hạn file size để tránh OOM."""
        size_mb = file_path.stat().st_size / (1024 * 1024)
        if size_mb > self.MAX_FILE_SIZE_MB:
            raise ValueError(
                f"File {file_path.name} ({size_mb:.1f}MB) vượt quá giới hạn {self.MAX_FILE_SIZE_MB}MB"
            )

    def _read_pdf(self, path: Path) -> str:
        """Đọc PDF bằng pypdf — nhẹ, không cần llama-index."""
        try:
            from pypdf import PdfReader
        except ImportError:
            raise ImportError("Cài đặt pypdf: pip install pypdf")
        
        reader = PdfReader(str(path))
        texts = []
        for page in reader.pages:
            text = page.extract_text()
            if text:
                texts.append(text)
        return "\n".join(texts)

    def _read_docx(self, path: Path) -> str:
        """Đọc Word bằng python-docx."""
        try:
            from docx import Document
        except ImportError:
            raise ImportError("Cài đặt python-docx: pip install python-docx")
        
        doc = Document(str(path))
        return "\n".join(p.text for p in doc.paragraphs if p.text.strip())

    def _read_xlsx(self, path: Path) -> str:
        """Đọc Excel bằng openpyxl."""
        try:
            from openpyxl import load_workbook
        except ImportError:
            raise ImportError("Cài đặt openpyxl: pip install openpyxl")
        
        wb = load_workbook(str(path), data_only=True)
        texts = []
        for sheet in wb.worksheets:
            for row in sheet.iter_rows(values_only=True):
                row_text = " | ".join(str(cell) for cell in row if cell is not None)
                if row_text.strip():
                    texts.append(row_text)
        return "\n".join(texts)

    def _read_txt(self, path: Path) -> str:
        """Đọc text file với encoding detection đơn giản."""
        for encoding in ["utf-8", "utf-16", "cp1252", "iso-8859-1"]:
            try:
                return path.read_text(encoding=encoding)
            except UnicodeDecodeError:
                continue
        raise ValueError(f"Không thể decode file: {path}")

    def _read_file_by_type(self, path: Path) -> str:
        """Router đọc file theo định dạng."""
        suffix = path.suffix.lower()
        readers = {
            ".pdf": self._read_pdf,
            ".docx": self._read_docx,
            ".xlsx": self._read_xlsx,
            ".txt": self._read_txt,
            ".csv": self._read_txt,
            ".md": self._read_txt,
        }
        reader = readers.get(suffix)
        if not reader:
            raise ValueError(f"Không có reader cho định dạng: {suffix}")
        return reader(path)

    def load_file(self, path: str) -> List[LoadedDocument]:
        """Đọc file đơn lẻ với kiểm tra nghiêm ngặt."""
        file_path = Path(path)

        if not file_path.exists():
            raise FileNotFoundError(f"Không tìm thấy file: {path}")

        if file_path.suffix.lower() not in self.SUPPORTED_EXTENSIONS:
            raise ValueError(
                f"Định dạng '{file_path.suffix}' không được hỗ trợ. "
                f"Chấp nhận: {', '.join(self.SUPPORTED_EXTENSIONS)}"
            )

        self._check_file_size(file_path)

        raw_text = self._read_file_by_type(file_path)
        cleaned = self._clean_text(raw_text)

        if not cleaned:
            raise ValueError(f"Không trích xuất được nội dung từ: {file_path.name}")

        return [LoadedDocument(
            text=cleaned,
            metadata={
                "file_name": file_path.name,
                "file_extension": file_path.suffix.lower().lstrip("."),
                "file_size_bytes": file_path.stat().st_size,
                "absolute_path": str(file_path.resolve()),
                "document_type": "file",
            },
        )]

    # ── Directory Loading ─────────────────────────────────────────────────────

    def _should_skip(self, path: Path) -> bool:
        """Kiểm tra path có nên bỏ qua không (hidden, venv, etc.)."""
        name = path.name
        skip_patterns = (".", "~", "node_modules", ".venv", "__pycache__", ".git")
        return any(pattern in str(path) for pattern in skip_patterns)

    def _iter_files(self, dir_path: Path) -> Iterator[Path]:
        for file_path in dir_path.rglob("*"):
            if file_path.is_file() and file_path.suffix.lower() in self.SUPPORTED_EXTENSIONS:
                yield file_path

    def load_directory(self, path: str) -> List[LoadedDocument]:
        """Quét đệ quy thư mục, stream từng file."""
        dir_path = Path(path)
        if not dir_path.is_dir():
            raise NotADirectoryError(f"Không phải thư mục: {path}")

        processed_docs: List[LoadedDocument] = []
        
        for file_path in self._iter_files(dir_path):
            try:
                self._check_file_size(file_path)
                raw_text = self._read_file_by_type(file_path)
                cleaned = self._clean_text(raw_text)
                
                if cleaned:
                    processed_docs.append(LoadedDocument(
                        text=cleaned,
                        metadata={
                            "file_name": file_path.name,
                            "file_extension": file_path.suffix.lower().lstrip("."),
                            "file_size_bytes": file_path.stat().st_size,
                            "absolute_path": str(file_path.resolve()),
                            "document_type": "file",
                        },
                    ))
            except Exception as exc:
                logger.warning("Bỏ qua file lỗi %s: %s", file_path, exc)
                continue

        logger.info("[Loader] Đã nạp %s tài liệu từ %s", len(processed_docs), path)
        return processed_docs